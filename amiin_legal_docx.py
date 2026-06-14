# =============================================================================
# AMIIN LEGAL — GÉNÉRATEUR DOCX
# Fichier  : amiin_legal_docx.py
# Version  : 2.0 — Améliorations : lettre client en section dédiée,
#                                   page de validation senior/junior,
#                                   alerte prescription en en-tête
# Dépend de: python-docx
# Install  : & "D:\projet\Agent Amiin\amiin_env\Scripts\python.exe" -m pip install python-docx
# =============================================================================

from __future__ import annotations

import logging
import re
from datetime import datetime
from typing import Optional

logger = logging.getLogger("amiin_legal_docx")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False
    logger.warning(
        "python-docx non installé.\n"
        '& "D:\\projet\\Agent Amiin\\amiin_env\\Scripts\\python.exe" -m pip install python-docx'
    )


# =============================================================================
# PALETTE
# =============================================================================

C_MARINE      = RGBColor(0x1A, 0x1A, 0x2E)
C_BLEU        = RGBColor(0x16, 0x21, 0x3E)
C_BLEU_ROYAL  = RGBColor(0x0F, 0x3A, 0x6E)
C_ROUGE       = RGBColor(0xC0, 0x39, 0x2B)
C_ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
C_VERT        = RGBColor(0x27, 0xAE, 0x60)
C_CORPS       = RGBColor(0x2C, 0x2C, 0x2C)
C_CITATION    = RGBColor(0x1A, 0x53, 0x76)
C_GRIS        = RGBColor(0x88, 0x88, 0x88)
C_LETTRE      = RGBColor(0x1A, 0x53, 0x76)   # Bleu pour la lettre client


# =============================================================================
# UTILITAIRES DOCX
# =============================================================================

def _font(run, name: str = "Garamond", size: int = 11,
          bold: bool = False, italic: bool = False,
          color: Optional[RGBColor] = None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _para(doc, texte: str, bold: bool = False, italic: bool = False,
          taille: int = 11, couleur: Optional[RGBColor] = None,
          alignement=None):
    para = doc.add_paragraph()
    if alignement:
        para.alignment = alignement
    run = para.add_run(texte)
    _font(run, size=taille, bold=bold, italic=italic, color=couleur)
    return para


def _hr(doc):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A1A2E")
    pBdr.append(bot)
    pPr.append(pBdr)


def _shading(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _tableau_markdown(doc, lignes: list[str]):
    lignes_data = [l for l in lignes if not re.match(r"^\|[-| :]+\|$", l.strip())]
    if len(lignes_data) < 2:
        return

    def cellules(ligne: str) -> list[str]:
        return [c.strip() for c in ligne.strip().strip("|").split("|")]

    header  = cellules(lignes_data[0])
    nb_cols = len(header)
    nb_rows = len(lignes_data)
    table   = doc.add_table(rows=nb_rows, cols=nb_cols)
    table.style = "Table Grid"

    for j, cel in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(cel)
        _font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shading(cell, "1A1A2E")

    for i, ligne in enumerate(lignes_data[1:], 1):
        cels = cellules(ligne)
        for j in range(min(len(cels), nb_cols)):
            cell  = table.rows[i].cells[j]
            cell.text = ""
            texte = cels[j]
            couleur = C_CORPS
            if "✅" in texte or "🟢" in texte:
                couleur = C_VERT
            elif "❌" in texte or "🔴" in texte:
                couleur = C_ROUGE
            elif "⚠️" in texte or "🟡" in texte:
                couleur = C_ORANGE
            run = cell.paragraphs[0].add_run(texte)
            _font(run, size=9, color=couleur)

    doc.add_paragraph()


def _inline(para, texte: str, taille: int = 10):
    pattern = r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)"
    for partie in re.split(pattern, texte):
        if not partie:
            continue
        if partie.startswith("***") and partie.endswith("***"):
            run = para.add_run(partie[3:-3])
            _font(run, size=taille, bold=True, italic=True)
        elif partie.startswith("**") and partie.endswith("**"):
            run = para.add_run(partie[2:-2])
            _font(run, size=taille, bold=True)
        elif partie.startswith("*") and partie.endswith("*") and len(partie) > 2:
            run = para.add_run(partie[1:-1])
            _font(run, size=taille, italic=True)
        elif partie.startswith("`") and partie.endswith("`"):
            run = para.add_run(partie[1:-1])
            _font(run, name="Courier New", size=taille - 1, color=C_CITATION)
        else:
            run = para.add_run(partie)
            _font(run, size=taille)


# =============================================================================
# EN-TÊTE DOCUMENT
# =============================================================================

def _entete(doc, type_affaire: str, dossier_id: str, date_generation: str,
            alerte_prescription: Optional[dict] = None):
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("AMIIN LEGAL")
    _font(run, size=18, bold=True, color=C_MARINE)

    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sous.add_run("Assistant Juridique IA — Cabinet d'Avocats de Djibouti")
    _font(run2, size=10, italic=True, color=C_CORPS)

    _hr(doc)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = info.add_run(f"Dossier n° {dossier_id}  |  {type_affaire}  |  {date_generation}")
    _font(run3, size=9, italic=True, color=C_CORPS)

    # AMÉLIORATION 1 — Alerte prescription dans l'en-tête si critique
    if alerte_prescription and alerte_prescription.get("niveau_alerte") in ("rouge", "expiree", "orange"):
        presc_para = doc.add_paragraph()
        presc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        couleur_presc = C_ROUGE if alerte_prescription["niveau_alerte"] in ("rouge", "expiree") else C_ORANGE
        run_p = presc_para.add_run(alerte_prescription.get("message", ""))
        _font(run_p, size=10, bold=True, color=couleur_presc)

    avert = doc.add_paragraph()
    avert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = avert.add_run(
        "⚠ Document de travail confidentiel — Validation obligatoire par l'avocat responsable"
    )
    _font(run4, size=9, bold=True, color=C_ROUGE)

    _hr(doc)
    doc.add_paragraph()


# =============================================================================
# AMÉLIORATION 3 — SECTION LETTRE CLIENT DÉDIÉE
# =============================================================================

def _section_lettre_client(doc, lettre_client: str):
    """
    Génère la section lettre client dans le DOCX avec un style distinct :
    - Fond encadré bleu discret
    - Typographie plus accessible (non juridique)
    - Mention de validation en rouge en haut
    - Séparation visuelle claire du reste du dossier
    """
    doc.add_page_break()

    # Titre de section
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titre.add_run("SECTION 7 — LETTRE DE SYNTHÈSE POUR LE CLIENT")
    _font(run_t, size=14, bold=True, color=C_BLEU)

    _hr(doc)

    # Avertissement avocat
    avert = doc.add_paragraph()
    avert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = avert.add_run(
        "À RELIRE ET SIGNER PAR L'AVOCAT RESPONSABLE AVANT ENVOI AU CLIENT"
    )
    _font(run_a, size=10, bold=True, color=C_ROUGE)

    doc.add_paragraph()

    # Note d'usage
    note = doc.add_paragraph()
    run_n = note.add_run(
        "Cette lettre est rédigée en français courant, sans jargon juridique, "
        "pour être directement compréhensible par le client. "
        "L'avocat la relit, la complète si nécessaire, et la signe avant envoi."
    )
    _font(run_n, size=9, italic=True, color=C_GRIS)

    doc.add_paragraph()
    _hr(doc)
    doc.add_paragraph()

    # Corps de la lettre
    for ligne in lettre_client.split("\n"):
        ligne_stripped = ligne.strip()

        if not ligne_stripped:
            doc.add_paragraph()
            continue

        if ligne_stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(ligne_stripped[2:])
            _font(run, size=13, bold=True, color=C_LETTRE)
            continue

        # Détection salutations / formule de politesse
        if any(ligne_stripped.startswith(s) for s in
               ["Madame", "Monsieur", "Cher", "Chère", "Cordialement", "Bien cordialement", "Sincèrement"]):
            p = doc.add_paragraph()
            run = p.add_run(ligne_stripped)
            _font(run, size=11, bold=True, color=C_CORPS)
            continue

        # Signature
        if ligne_stripped.startswith("Maître") or ligne_stripped.startswith("[Signature"):
            p = doc.add_paragraph()
            run = p.add_run(ligne_stripped)
            _font(run, size=11, italic=True, color=C_CORPS)
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        _inline(p, ligne_stripped, taille=11)

    doc.add_paragraph()
    _hr(doc)


# =============================================================================
# AMÉLIORATION 2 — PAGE DE VALIDATION SENIOR/JUNIOR
# =============================================================================

def _page_validation(doc, dossier_id: str, date_generation: str,
                     collaborateur: Optional[str] = None,
                     annotation_senior: Optional[str] = None,
                     statut_validation: str = "en_attente"):
    """
    Génère une page de validation à la fin du DOCX :
    - Bloc "Préparé par le collaborateur" (signature junior)
    - Bloc "Validation du senior" (zone de signature + annotation)
    - Checklist de validation
    Cette page est à imprimer et signer physiquement si nécessaire,
    ou à compléter dans l'interface Amiin Legal.
    """
    doc.add_page_break()

    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titre.add_run("FICHE DE VALIDATION — USAGE INTERNE CABINET")
    _font(run_t, size=13, bold=True, color=C_MARINE)

    _hr(doc)
    doc.add_paragraph()

    # Informations du dossier
    _para(doc, f"Dossier n° {dossier_id}  |  Généré le {date_generation}",
          italic=True, taille=9, couleur=C_GRIS)
    doc.add_paragraph()

    # Bloc collaborateur
    _para(doc, "1. PRÉPARÉ PAR", bold=True, taille=11, couleur=C_BLEU)
    table_collab = doc.add_table(rows=3, cols=2)
    table_collab.style = "Table Grid"
    lignes_c = [
        ("Collaborateur",   collaborateur or "[À COMPLÉTER]"),
        ("Date de remise",  date_generation),
        ("Signature",       ""),
    ]
    for i, (label, valeur) in enumerate(lignes_c):
        table_collab.rows[i].cells[0].text = label
        table_collab.rows[i].cells[1].text = valeur
        run_l = table_collab.rows[i].cells[0].paragraphs[0].runs
        if run_l:
            _font(run_l[0], size=10, bold=True)
        run_v = table_collab.rows[i].cells[1].paragraphs[0].runs
        if run_v:
            _font(run_v[0], size=10)

    doc.add_paragraph()

    # Checklist validation
    _para(doc, "2. POINTS DE VALIDATION (à cocher par le senior)", bold=True, taille=11, couleur=C_BLEU)
    checklist_items = [
        "Qualification juridique vérifiée et correcte",
        "Articles de loi cités vérifiés dans les codes djiboutiens",
        "Stratégie retenue adaptée aux faits et aux intérêts du client",
        "Prétentions chiffrées vérifiées et validées avec le client",
        "Délais procéduraux vérifiés — aucune forclusion en cours",
        "Actes rédigés relus et corrigés si nécessaire",
        "Lettre client (Section 7) relue et prête à signature",
        "Aucun conflit d'intérêts identifié",
    ]
    for item in checklist_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"☐  {item}")
        _font(run, size=10, color=C_CORPS)

    doc.add_paragraph()

    # Annotation senior
    _para(doc, "3. ANNOTATIONS DU SENIOR", bold=True, taille=11, couleur=C_BLEU)
    if annotation_senior:
        p = doc.add_paragraph()
        run_ann = p.add_run(annotation_senior)
        _font(run_ann, size=10, italic=True, color=C_CORPS)
    else:
        # Zone vide pour annotations manuscrites
        for _ in range(5):
            p = doc.add_paragraph()
            run_ligne = p.add_run("_" * 80)
            _font(run_ligne, size=10, color=C_GRIS)

    doc.add_paragraph()

    # Décision finale
    _para(doc, "4. DÉCISION DU SENIOR", bold=True, taille=11, couleur=C_BLEU)
    statut_map = {
        "en_attente": ("EN ATTENTE", C_ORANGE),
        "valide":     ("✅ VALIDÉ", C_VERT),
        "renvoye":    ("↩ RENVOYÉ POUR CORRECTION", C_ROUGE),
    }
    label_statut, couleur_statut = statut_map.get(statut_validation, ("EN ATTENTE", C_ORANGE))

    p_statut = doc.add_paragraph()
    run_s = p_statut.add_run(f"Statut : {label_statut}")
    _font(run_s, size=12, bold=True, color=couleur_statut)

    doc.add_paragraph()

    table_sign = doc.add_table(rows=2, cols=2)
    table_sign.style = "Table Grid"
    headers = ["Avocat senior validateur", "Date de validation"]
    for j, h in enumerate(headers):
        cell = table_sign.rows[0].cells[j]
        cell.text = h
        if cell.paragraphs[0].runs:
            _font(cell.paragraphs[0].runs[0], size=10, bold=True)
        table_sign.rows[1].cells[j].text = ""

    _para(doc, "\nSignature :", taille=10, couleur=C_CORPS)
    for _ in range(3):
        doc.add_paragraph()


# =============================================================================
# PARSEUR MARKDOWN → DOCX (corps principal)
# =============================================================================

def _parser_corps(doc, markdown: str):
    """
    Parse le Markdown principal (hors lettre client et page validation).
    Gère H1-H4, tableaux, listes, citations, actes, alertes.
    """
    lignes         = markdown.split("\n")
    i              = 0
    dans_tableau   = False
    buf_tableau:   list[str] = []
    dans_acte      = False
    num_acte       = 0

    while i < len(lignes):
        ligne = lignes[i]

        # ── Marqueurs lettre client — SAUTER le bloc complet ─────────────────
        # (la lettre client est générée séparément via _section_lettre_client)
        if "<!-- DEBUT_LETTRE_CLIENT -->" in ligne:
            while i < len(lignes) and "<!-- FIN_LETTRE_CLIENT -->" not in lignes[i]:
                i += 1
            i += 1
            continue

        # ── Marqueurs actes ───────────────────────────────────────────────────
        if "<!-- DEBUT_DOCUMENT_" in ligne:
            dans_acte = True
            num_acte += 1
            if buf_tableau:
                _tableau_markdown(doc, buf_tableau)
                buf_tableau = []
                dans_tableau = False
            doc.add_page_break()
            _para(doc, f"ANNEXE {chr(64 + num_acte)} — ACTE RÉDIGÉ",
                  bold=True, taille=13, couleur=C_MARINE)
            _hr(doc)
            i += 1
            continue

        if "<!-- FIN_DOCUMENT_" in ligne:
            dans_acte = False
            _hr(doc)
            i += 1
            continue

        # ── Tableaux ──────────────────────────────────────────────────────────
        if ligne.strip().startswith("|"):
            if not dans_tableau:
                dans_tableau = True
                buf_tableau  = []
            buf_tableau.append(ligne)
            i += 1
            continue
        else:
            if dans_tableau:
                _tableau_markdown(doc, buf_tableau)
                dans_tableau = False
                buf_tableau  = []

        # ── Titres ────────────────────────────────────────────────────────────
        if ligne.startswith("#### "):
            _para(doc, ligne[5:].strip(), bold=True, taille=10, couleur=C_BLEU_ROYAL)
            i += 1
            continue
        if ligne.startswith("### "):
            _para(doc, ligne[4:].strip(), bold=True, taille=11, couleur=C_BLEU_ROYAL)
            i += 1
            continue
        if ligne.startswith("## "):
            doc.add_paragraph()
            _para(doc, ligne[3:].strip(), bold=True, taille=13, couleur=C_BLEU)
            _hr(doc)
            i += 1
            continue
        if ligne.startswith("# "):
            doc.add_paragraph()
            _para(doc, ligne[2:].strip(), bold=True, taille=16, couleur=C_MARINE,
                  alignement=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # ── Séparateurs ───────────────────────────────────────────────────────
        if ligne.strip() in ("---", "***", "___"):
            _hr(doc)
            i += 1
            continue

        # ── Listes ───────────────────────────────────────────────────────────
        if re.match(r"^[-*] ", ligne):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, ligne[2:].strip(), taille=10)
            i += 1
            continue
        if re.match(r"^\d+\. ", ligne):
            p = doc.add_paragraph(style="List Number")
            _inline(p, re.sub(r"^\d+\. ", "", ligne).strip(), taille=10)
            i += 1
            continue
        if re.match(r"^- \[.?\] ", ligne):
            texte = re.sub(r"^- \[.?\] ", "", ligne).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            checked = "✅" if "[x]" in ligne or "[X]" in ligne else "☐"
            run = p.add_run(f"{checked}  {texte}")
            _font(run, size=10, color=C_CORPS)
            i += 1
            continue

        # ── Citations > ───────────────────────────────────────────────────────
        if ligne.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(ligne[2:].strip())
            _font(run, size=10, italic=True, color=C_CITATION)
            i += 1
            continue

        # ── Alertes ───────────────────────────────────────────────────────────
        if ligne.strip().startswith("🔴"):
            p = doc.add_paragraph()
            run = p.add_run(ligne.strip())
            _font(run, size=10, bold=True, color=C_ROUGE)
            i += 1
            continue
        if ligne.strip().startswith("⚠️"):
            p = doc.add_paragraph()
            run = p.add_run(ligne.strip())
            _font(run, size=10, bold=True, color=C_ORANGE)
            i += 1
            continue
        if ligne.strip().startswith("✅"):
            p = doc.add_paragraph()
            run = p.add_run(ligne.strip())
            _font(run, size=10, bold=True, color=C_VERT)
            i += 1
            continue

        # ── Paragraphe vide ───────────────────────────────────────────────────
        if not ligne.strip():
            doc.add_paragraph()
            i += 1
            continue

        # ── Paragraphe normal ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        _inline(p, ligne, taille=10 if dans_acte else 10.5)
        i += 1

    if buf_tableau:
        _tableau_markdown(doc, buf_tableau)


# =============================================================================
# PIED DE PAGE
# =============================================================================

def _pied_page(doc, dossier_id: str):
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        f"AMIIN LEGAL v2.0  |  Dossier {dossier_id}  |  "
        "Document de travail — validation obligatoire avant usage"
    )
    _font(run, size=8, italic=True, color=C_CORPS)


# =============================================================================
# FONCTION PRINCIPALE
# =============================================================================

def generer_docx_dossier(
    dossier_markdown:    str,
    output_path:         str,
    type_affaire:        str,
    dossier_id:          str,
    date_generation:     str = "",
    lettre_client:       Optional[str] = None,
    alerte_prescription: Optional[dict] = None,
    collaborateur:       Optional[str] = None,
    annotation_senior:   Optional[str] = None,
    statut_validation:   str = "en_attente",
) -> str:
    """
    Convertit un dossier juridique Markdown en fichier DOCX professionnel.

    Args:
        dossier_markdown    : Texte Markdown complet généré par Claude
        output_path         : Chemin de sortie .docx
        type_affaire        : Label de l'affaire
        dossier_id          : Identifiant court (ex: "A3F7B2")
        date_generation     : Date formatée
        lettre_client       : AMÉLIORATION 3 — Texte lettre client extrait
        alerte_prescription : AMÉLIORATION 1 — Dict alerte prescription
        collaborateur       : AMÉLIORATION 2 — Nom du collaborateur
        annotation_senior   : AMÉLIORATION 2 — Annotation du senior
        statut_validation   : AMÉLIORATION 2 — "en_attente"|"valide"|"renvoye"

    Returns:
        output_path si succès.
    """
    if not DOCX_DISPONIBLE:
        raise RuntimeError(
            "python-docx non installé.\n"
            '& "D:\\projet\\Agent Amiin\\amiin_env\\Scripts\\python.exe" -m pip install python-docx'
        )

    if not date_generation:
        date_generation = datetime.now().strftime("%d/%m/%Y à %Hh%M")

    logger.info(f"[{dossier_id}] Génération DOCX → {output_path}")

    doc = Document()

    # ── Paramètres de page A4 ─────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── En-tête (avec alerte prescription si critique) ────────────────────────
    _entete(doc, type_affaire, dossier_id, date_generation, alerte_prescription)

    # ── Corps principal (sections 1-6 + 8-11) ────────────────────────────────
    _parser_corps(doc, dossier_markdown)

    # ── AMÉLIORATION 3 — Section lettre client dédiée ─────────────────────────
    if lettre_client:
        _section_lettre_client(doc, lettre_client)
    else:
        # Fallback : extraire depuis le markdown si non fournie
        debut = "<!-- DEBUT_LETTRE_CLIENT -->"
        fin   = "<!-- FIN_LETTRE_CLIENT -->"
        if debut in dossier_markdown and fin in dossier_markdown:
            start = dossier_markdown.index(debut) + len(debut)
            end   = dossier_markdown.index(fin)
            lettre_extraite = dossier_markdown[start:end].strip()
            _section_lettre_client(doc, lettre_extraite)

    # ── AMÉLIORATION 2 — Page de validation ──────────────────────────────────
    _page_validation(
        doc,
        dossier_id         = dossier_id,
        date_generation    = date_generation,
        collaborateur      = collaborateur,
        annotation_senior  = annotation_senior,
        statut_validation  = statut_validation,
    )

    # ── Pied de page ──────────────────────────────────────────────────────────
    _pied_page(doc, dossier_id)

    # ── Sauvegarde ────────────────────────────────────────────────────────────
    doc.save(output_path)
    logger.info(f"[{dossier_id}] DOCX sauvegardé : {output_path}")
    return output_path


# =============================================================================
# TEST STANDALONE
# =============================================================================
# & "D:\projet\Agent Amiin\amiin_env\Scripts\python.exe" amiin_legal_docx.py

if __name__ == "__main__":
    import sys

    LETTRE_TEST = """
Djibouti, le 06 janvier 2025

Madame, Monsieur Hassan Ahmed,

Nous avons bien reçu votre dossier concernant votre licenciement par la Société DJIB IMPORT SARL.

Votre situation en quelques mots : Votre licenciement présente des irrégularités importantes. L'employeur n'a pas respecté la procédure légale et n'a pas fourni de motif valable. Vous êtes dans une position favorable pour obtenir réparation.

Ce que nous allons faire : Nous allons déposer une requête devant le Tribunal du Travail de Djibouti pour obtenir le paiement de vos indemnités et des dommages-intérêts. Cette procédure dure en moyenne 3 à 6 mois.

Ce que vous devez nous fournir :
- Votre contrat de travail
- Vos 3 derniers bulletins de salaire
- Toutes les lettres reçues de l'employeur
- Votre attestation CNSS

Prochaine étape : Merci de nous apporter ces documents avant le 15 janvier 2025.

Cordialement,

Maître [NOM DE L'AVOCAT]
Avocat au Barreau de Djibouti
"""

    MARKDOWN_TEST = """
# DOSSIER AMIIN LEGAL
## Contentieux du Travail — Licenciement Abusif — Hassan Ahmed

---

## SECTION 1 — RÉSUMÉ EXÉCUTIF

M. Hassan Ahmed, salarié depuis 7 ans, a été licencié sans motif réel et sérieux.
Le dossier est juridiquement solide. Action prud'homale recommandée.

🔴 DÉLAI CRITIQUE : Dépôt requête Tribunal du Travail avant le 12/01/2025 (dans 6 jours)

---

## SECTION 2 — IDENTIFICATION DES PARTIES

| Rôle | Identité | Qualité juridique | Domicile |
|------|----------|-------------------|----------|
| Demandeur | M. Hassan Ahmed | Personne physique | Djibouti-Ville |
| Défendeur | Société DJIB IMPORT SARL | Personne morale | Zone Industrielle |

---

## SECTION 3 — QUALIFICATION JURIDIQUE

**Qualification retenue** : Licenciement abusif sans cause réelle et sérieuse
**Fondement principal** : [CTD, Art. licenciement]

⚠️ POINT À VÉRIFIER : Numéro exact de l'article sur le licenciement dans le CTD

---

## SECTION 10 — CHIFFRAGE

| Poste | Base de calcul | Montant FDJ | Fondement |
|-------|----------------|-------------|-----------|
| Indemnité de licenciement | 7 ans × salaire | 350 000 FDJ | CTD |
| Préavis | 3 mois | 150 000 FDJ | CTD |
| **TOTAL** | | **500 000 FDJ** | |

---

## SECTION 11 — ACTES RÉDIGÉS

### ANNEXE A — MISE EN DEMEURE

<!-- DEBUT_DOCUMENT_A -->

Maître [NOM AVOCAT]
Avocat au Barreau de Djibouti

Djibouti, le 06 janvier 2025

**OBJET : Mise en demeure — Licenciement abusif de M. Hassan Ahmed**

Par la présente, nous mettons en demeure la Société DJIB IMPORT SARL de régulariser la situation.

*Signature et sceau de Maître [NOM]*

<!-- FIN_DOCUMENT_A -->
"""

    ALERTE_TEST = {
        "niveau_alerte":  "rouge",
        "message":        "🔴 DÉLAI CRITIQUE : Dépôt requête avant le 12/01/2025 (dans 6 jours) [CTD, Art. prescription]",
        "delai_label":    "3 ans",
        "action":         "Saisine Tribunal du Travail",
        "article":        "CTD, Art. prescription",
        "jours_ecoules":  779,
        "jours_restants": 6,
        "date_limite":    "12/01/2025",
    }

    test_path = "./test_dossier_v2.docx"
    try:
        generer_docx_dossier(
            dossier_markdown    = MARKDOWN_TEST,
            output_path         = test_path,
            type_affaire        = "Contentieux du Travail — Licenciement Abusif",
            dossier_id          = "TEST02",
            date_generation     = datetime.now().strftime("%d/%m/%Y à %Hh%M"),
            lettre_client       = LETTRE_TEST,
            alerte_prescription = ALERTE_TEST,
            collaborateur       = "M. Ali Ibrahim",
            annotation_senior   = "Vérifier l'article exact sur le préavis. Montant D&I à revoir.",
            statut_validation   = "renvoye",
        )
        print(f"✅ DOCX v2 généré avec succès : {test_path}")
    except Exception as e:
        print(f"❌ Erreur : {e}", file=sys.stderr)
        sys.exit(1)
